import sys
from PySide6.QtWidgets import QApplication, QFileDialog, QColorDialog, QMessageBox, QInputDialog
from PySide6.QtCore import QTimer
from PySide6.QtGui import QTextCharFormat, QColor
import markdown
import json
from database import Database, NoteCategory
from ui.main_window import MainWindow

class NoteTypewriter:
    def __init__(self):
        self.db = Database()
        self.window = MainWindow()
        
        # Connect signals
        self.window.note_selected.connect(self.load_note)
        self.window.note_deleted.connect(self.delete_note)
        self.window.search_changed.connect(self.search_notes)
        self.window.category_changed.connect(self.change_category)
        
        # Connect menu action signals
        self.window.new_note_requested.connect(self.new_note)
        self.window.save_note_requested.connect(self.save_note)
        self.window.export_requested.connect(self.export_note)
        
        # Set up auto-save timer
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(30000)  # Auto-save every 30 seconds
        
        # Initialize state
        self.current_note = None
        self.current_category = NoteCategory.ALL
        
        # Load existing notes
        self.refresh_notes()
    
    def refresh_notes(self):
        notes = self.db.get_all_notes(category=self.current_category)
        self.window.refresh_note_list(notes)
    
    def new_note(self):
        title, ok = QInputDialog.getText(self.window, "New Note", "Enter note title:")
        if ok and title:
            note = self.db.create_note(
                title=title,
                category=self.current_category
            )
            self.refresh_notes()
            self.load_note(note.id)
    
    def load_note(self, note_id):
        note = self.db.get_note(note_id)
        if note:
            self.current_note = note
            self.window.set_note_content(
                note.title,
                note.content,
                note.tags,
                metadata={
                    'created_at': note.created_at,
                    'updated_at': note.updated_at
                }
            )
    
    def save_note(self):
        if self.current_note:
            content = self.window.get_note_content()
            html_content = self.window.editor_widget.preview.md.convert(content)
            self.db.update_note(
                self.current_note.id,
                content=content,
                html_content=html_content,
                tags=self.window.get_note_tags()
            )
            self.refresh_notes()
    
    def auto_save(self):
        if self.current_note and self.window.editor.document().isModified():
            self.save_note()
    
    def delete_note(self, note_id):
        reply = QMessageBox.question(
            self.window, "Delete Note",
            "Are you sure you want to delete this note?",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.db.delete_note(note_id)
            self.current_note = None
            self.window.editor.clear()
            self.window.tags_input.clear()
            self.refresh_notes()
    
    def search_notes(self, query):
        notes = self.db.search_notes(query, category=self.current_category)
        self.window.refresh_note_list(notes)
    
    def change_category(self, category_name):
        try:
            self.current_category = NoteCategory(category_name)
        except ValueError:
            self.current_category = NoteCategory.ALL
        self.refresh_notes()
    
    def export_note(self, format_type):
        if not self.current_note:
            return
        
        file_filters = {
            "markdown": "Markdown Files (*.md)",
            "html": "HTML Files (*.html)",
            "pdf": "PDF Files (*.pdf)",
            "docx": "Word Document (*.docx)"
        }
        
        file_name, _ = QFileDialog.getSaveFileName(
            self.window,
            "Export Note",
            "",
            file_filters[format_type]
        )
        
        if file_name:
            note_data = self.db.export_note(self.current_note.id, format=format_type)
            
            if note_data:
                try:
                    if format_type == "markdown":
                        with open(file_name, 'w', encoding='utf-8') as f:
                            f.write(f"# {note_data['title']}\n\n")
                            f.write(note_data['content'])
                    elif format_type == "html":
                        html_content = self.window.editor_widget.preview.md.convert(note_data['content'])
                        with open(file_name, 'w', encoding='utf-8') as f:
                            f.write(f"<h1>{note_data['title']}</h1>\n")
                            f.write(html_content)
                    elif format_type == "pdf":
                        import pdfkit
                        html_content = self.window.editor_widget.preview.md.convert(note_data['content'])
                        styled_html = f"""
                        <style>
                            body {{
                                font-family: 'Arial', sans-serif;
                                line-height: 1.6;
                                margin: 40px;
                            }}
                            h1 {{ color: #ff69b4; }}
                            pre {{ background-color: #f5f5f5; padding: 10px; }}
                        </style>
                        <h1>{note_data['title']}</h1>
                        {html_content}
                        """
                        pdfkit.from_string(styled_html, file_name)
                    elif format_type == "docx":
                        from docx import Document
                        doc = Document()
                        doc.add_heading(note_data['title'], 0)
                        doc.add_paragraph(note_data['content'])
                        doc.save(file_name)
                    
                    QMessageBox.information(
                        self.window,
                        "Success",
                        f"Note exported successfully to {file_name}"
                    )
                except Exception as e:
                    QMessageBox.warning(
                        self.window,
                        "Error",
                        f"Failed to export note: {str(e)}"
                    )

def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    
    note_app = NoteTypewriter()
    note_app.window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main() 